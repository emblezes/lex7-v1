"""Export DOCX — generation de livrables en Word."""

import io
import logging
import re
from datetime import datetime

logger = logging.getLogger(__name__)


def export_livrable_docx(
    title: str,
    content_markdown: str,
    company_name: str,
    livrable_type: str = "",
    metadata: dict | None = None,
) -> bytes:
    """Genere un fichier DOCX a partir de contenu markdown.

    Args:
        title: Titre du document
        content_markdown: Contenu en markdown
        company_name: Nom du client
        livrable_type: Type de livrable
        metadata: Metadonnees (auteur, date, etc.)

    Returns:
        bytes du fichier DOCX
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("pip install python-docx pour le support DOCX")

    doc = Document()
    meta = metadata or {}

    # Style du document
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # En-tete
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(f"LegiX — {company_name}")
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(128, 128, 128)

    # Titre principal
    doc.add_heading(title, level=0)

    # Metadonnees
    meta_text = []
    if livrable_type:
        meta_text.append(f"Type : {livrable_type}")
    meta_text.append(f"Date : {meta.get('date', datetime.now().strftime('%d/%m/%Y'))}")
    if meta.get("auteur"):
        meta_text.append(f"Auteur : {meta['auteur']}")
    if meta.get("target_audience"):
        meta_text.append(f"Destinataire : {meta['target_audience']}")

    if meta_text:
        meta_para = doc.add_paragraph(" | ".join(meta_text))
        meta_para.runs[0].font.size = Pt(9)
        meta_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph("")  # Ligne vide

    # Convertir le markdown en paragraphes Word
    _markdown_to_docx(doc, content_markdown)

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"Document genere par LegiX pour {company_name}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(160, 160, 160)

    # Sauvegarder en bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _markdown_to_docx(doc, markdown: str):
    """Convertit du markdown simplifie en paragraphes Word."""
    from docx.shared import Pt, RGBColor

    lines = markdown.split("\n")
    in_list = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_list:
                in_list = False
            doc.add_paragraph("")
            continue

        # Titres
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)

        # Listes a puces
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, text)
            in_list = True

        # Listes numerotees
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, text)
            in_list = True

        # Blockquotes
        elif stripped.startswith("> "):
            para = doc.add_paragraph()
            run = para.add_run(stripped[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Separateurs
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 50)

        # Texte normal
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)


def _add_formatted_text(para, text: str):
    """Ajoute du texte avec formatage bold/italic."""
    from docx.shared import RGBColor

    # Regex pour bold et italic inline
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.color.rgb = RGBColor(180, 60, 60)
        else:
            para.add_run(part)
