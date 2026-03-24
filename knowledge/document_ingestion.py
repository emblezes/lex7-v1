"""Ingestion de documents client — parse, chunk, embed, store.

Pipeline :
1. Upload fichier (PDF, DOCX, TXT, MD) ou texte brut
2. Extraction du texte
3. Enrichissement IA (themes, positions, resume)
4. Chunking semantique
5. Embedding des chunks
6. Stockage en DB (ClientDocument + DocumentChunk)

Usage :
    from legix.knowledge.document_ingestion import ingest_document
    doc = await ingest_document(db, profile_id=1, file_path="rapport.pdf", doc_type="position_paper")
"""

import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from legix.core.models import ClientDocument

logger = logging.getLogger(__name__)


# ── Extraction texte par format ──────────────────────────────────

async def extract_text(file_path: str) -> str:
    """Extrait le texte d'un fichier selon son format."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8")

    elif suffix == ".pdf":
        return await _extract_pdf(path)

    elif suffix in (".docx", ".doc"):
        return await _extract_docx(path)

    elif suffix == ".html":
        return await _extract_html(path)

    else:
        raise ValueError(f"Format non supporte: {suffix}. Formats acceptes: .txt, .md, .pdf, .docx, .html")


async def _extract_pdf(path: Path) -> str:
    """Extrait le texte d'un PDF via pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pip install pypdf pour le support PDF")

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


async def _extract_docx(path: Path) -> str:
    """Extrait le texte d'un fichier DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("pip install python-docx pour le support DOCX")

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


async def _extract_html(path: Path) -> str:
    """Extrait le texte d'un fichier HTML."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("pip install beautifulsoup4 pour le support HTML")

    html = path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    # Supprimer scripts et styles
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


# ── Chunking semantique ─────────────────────────────────────────

def chunk_text(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    separator: str = "\n\n",
) -> list[dict]:
    """Decoupe le texte en chunks avec overlap.

    Essaie de couper aux frontieres semantiques (paragraphes, titres).
    Retourne une liste de dicts {text, start_idx, end_idx, chunk_idx}.
    """
    if not text or not text.strip():
        return []

    # D'abord, essayer de couper par paragraphes
    paragraphs = text.split(separator)
    if len(paragraphs) <= 1:
        # Fallback: couper par phrases
        paragraphs = re.split(r"(?<=[.!?])\s+", text)

    chunks = []
    current_chunk = ""
    current_start = 0
    chunk_idx = 0

    for para in paragraphs:
        if not para.strip():
            continue

        # Si ajouter ce paragraphe depasse la taille max
        if len(current_chunk) + len(para) > chunk_size and current_chunk:
            chunks.append({
                "text": current_chunk.strip(),
                "start_idx": current_start,
                "end_idx": current_start + len(current_chunk),
                "chunk_idx": chunk_idx,
            })
            chunk_idx += 1

            # Overlap : garder les derniers caracteres
            if chunk_overlap > 0 and len(current_chunk) > chunk_overlap:
                overlap_text = current_chunk[-chunk_overlap:]
                current_start = current_start + len(current_chunk) - chunk_overlap
                current_chunk = overlap_text + separator + para
            else:
                current_start = current_start + len(current_chunk)
                current_chunk = para
        else:
            if current_chunk:
                current_chunk += separator + para
            else:
                current_chunk = para

    # Dernier chunk
    if current_chunk.strip():
        chunks.append({
            "text": current_chunk.strip(),
            "start_idx": current_start,
            "end_idx": current_start + len(current_chunk),
            "chunk_idx": chunk_idx,
        })

    return chunks


# ── Enrichissement IA ────────────────────────────────────────────

async def enrich_document(text: str, doc_type: str) -> dict:
    """Analyse le document avec Claude pour en extraire :
    - themes : liste de themes reglementaires
    - key_positions : positions et arguments identifies
    - summary : resume en 3-5 phrases
    - mentioned_stakeholders : acteurs mentionnes
    """
    from legix.core.config import settings

    try:
        import anthropic
    except ImportError:
        logger.warning("anthropic non installe, enrichissement IA desactive")
        return {
            "themes": [],
            "key_positions": [],
            "summary": text[:500] + "...",
            "mentioned_stakeholders": [],
        }

    client = anthropic.AsyncAnthropic(api_key=settings.ANTHROPIC_API_KEY)

    prompt = f"""Analyse ce document de type "{doc_type}" et extrais :

1. THEMES : liste des themes reglementaires (max 5). Choisis parmi :
   sante, environnement/climat, fiscalite, numerique/tech, transport, energie,
   agriculture, education, defense, travail/emploi, logement, justice, finance,
   industrie, commerce, securite, culture, immigration, international

2. KEY_POSITIONS : les positions, arguments et propositions cles du document (max 5).
   Pour chaque position : {{"position": "...", "argument": "...", "force": "forte/moyenne/faible"}}

3. SUMMARY : resume factuel en 3-5 phrases.

4. STAKEHOLDERS : acteurs mentionnes (personnes, organisations, institutions).

Reponds en JSON strict :
{{"themes": [...], "key_positions": [...], "summary": "...", "mentioned_stakeholders": [...]}}

DOCUMENT (premiers 4000 caracteres) :
{text[:4000]}"""

    response = await client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )

    try:
        # Extraire le JSON de la reponse
        content = response.content[0].text
        # Chercher le JSON dans la reponse
        json_match = re.search(r"\{[\s\S]*\}", content)
        if json_match:
            return json.loads(json_match.group())
    except (json.JSONDecodeError, IndexError, AttributeError) as e:
        logger.warning("Enrichissement IA: parsing JSON echoue: %s", e)

    return {
        "themes": [],
        "key_positions": [],
        "summary": text[:500] + "...",
        "mentioned_stakeholders": [],
    }


# ── Embeddings ───────────────────────────────────────────────────

async def compute_embeddings(texts: list[str]) -> list[list[float]]:
    """Calcule les embeddings pour une liste de textes.

    Utilise l'API Voyager d'Anthropic ou un modele local.
    Fallback : embedding simple TF-IDF si pas d'API.
    """
    from legix.core.config import settings

    # Option 1 : Voyager embeddings (Anthropic)
    try:
        import httpx
        async with httpx.AsyncClient() as client:
            response = await client.post(
                "https://api.voyageai.com/v1/embeddings",
                headers={"Authorization": f"Bearer {settings.VOYAGE_API_KEY}"},
                json={
                    "model": "voyage-3",
                    "input": texts,
                    "input_type": "document",
                },
                timeout=30.0,
            )
            if response.status_code == 200:
                data = response.json()
                return [item["embedding"] for item in data["data"]]
    except Exception as e:
        logger.debug("Voyage embeddings indisponible: %s", e)

    # Option 2 : sentence-transformers local
    try:
        from sentence_transformers import SentenceTransformer
        model = SentenceTransformer("all-MiniLM-L6-v2")
        embeddings = model.encode(texts)
        return [emb.tolist() for emb in embeddings]
    except ImportError:
        pass

    # Option 3 : fallback simple (TF-IDF hash)
    logger.warning("Aucun modele d'embedding disponible, utilisation du fallback hash")
    return [_simple_hash_embedding(t) for t in texts]


def _simple_hash_embedding(text: str, dim: int = 384) -> list[float]:
    """Embedding de fallback par hashing — pour dev uniquement."""
    import hashlib
    import struct
    h = hashlib.sha512(text.encode()).digest()
    # Etendre le hash pour avoir assez de dimensions
    extended = h * (dim // 64 + 1)
    values = []
    for i in range(dim):
        byte_val = extended[i]
        values.append((byte_val / 255.0) * 2 - 1)  # Normaliser entre -1 et 1
    return values


# ── Pipeline principal ───────────────────────────────────────────

async def ingest_document(
    db: AsyncSession,
    profile_id: int,
    doc_type: str,
    title: str,
    content: Optional[str] = None,
    file_path: Optional[str] = None,
    file_name: Optional[str] = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> ClientDocument:
    """Pipeline complet d'ingestion d'un document client.

    Args:
        db: Session DB async
        profile_id: ID du profil client
        doc_type: Type (position_paper, internal_note, email, communication, rapport, presentation)
        title: Titre du document
        content: Texte brut (si pas de fichier)
        file_path: Chemin du fichier a ingerer
        file_name: Nom original du fichier
        chunk_size: Taille max d'un chunk (caracteres)
        chunk_overlap: Overlap entre chunks

    Returns:
        ClientDocument cree et enrichi
    """
    # 1. Extraire le texte
    if content:
        text = content
    elif file_path:
        text = await extract_text(file_path)
        if not file_name:
            file_name = Path(file_path).name
    else:
        raise ValueError("content ou file_path requis")

    if not text.strip():
        raise ValueError("Document vide apres extraction")

    logger.info("Ingestion: %s (%d chars) pour profile #%d", title, len(text), profile_id)

    # 2. Enrichissement IA
    enrichment = await enrich_document(text, doc_type)

    # 3. Creer le ClientDocument
    doc = ClientDocument(
        profile_id=profile_id,
        doc_type=doc_type,
        title=title,
        content=text,
        file_path=file_path,
        file_name=file_name,
        themes=json.dumps(enrichment.get("themes", []), ensure_ascii=False),
        key_positions=json.dumps(enrichment.get("key_positions", []), ensure_ascii=False),
        mentioned_stakeholders=json.dumps(
            enrichment.get("mentioned_stakeholders", []), ensure_ascii=False
        ),
        summary=enrichment.get("summary", ""),
    )
    db.add(doc)
    await db.flush()  # Pour avoir doc.id

    # 4. Chunking
    chunks = chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    logger.info("  -> %d chunks generes", len(chunks))

    # 5. Embeddings
    if chunks:
        chunk_texts = [c["text"] for c in chunks]
        embeddings = await compute_embeddings(chunk_texts)

        # 6. Stocker les chunks
        from legix.knowledge.models import DocumentChunk
        for chunk_data, embedding in zip(chunks, embeddings):
            chunk = DocumentChunk(
                document_id=doc.id,
                profile_id=profile_id,
                chunk_idx=chunk_data["chunk_idx"],
                content=chunk_data["text"],
                embedding=json.dumps(embedding),  # JSON pour SQLite, vector pour pgvector
                start_idx=chunk_data["start_idx"],
                end_idx=chunk_data["end_idx"],
            )
            db.add(chunk)

    await db.commit()
    await db.refresh(doc)

    logger.info(
        "Document ingere: #%d '%s' (%d chunks, themes=%s)",
        doc.id, title, len(chunks), enrichment.get("themes", []),
    )
    return doc


async def search_knowledge_base(
    db: AsyncSession,
    profile_id: int,
    query: str,
    top_k: int = 5,
    theme_filter: Optional[str] = None,
) -> list[dict]:
    """Recherche semantique dans la knowledge base du client.

    Utilise la similarite cosinus entre l'embedding de la query
    et les embeddings des chunks stockes.

    Args:
        db: Session DB
        profile_id: ID du profil client
        query: Question ou texte de recherche
        top_k: Nombre de resultats a retourner
        theme_filter: Filtrer par theme (optionnel)

    Returns:
        Liste de dicts {document_title, chunk_text, score, doc_type, themes}
    """
    from legix.knowledge.models import DocumentChunk

    # Embedding de la query
    query_embedding = (await compute_embeddings([query]))[0]

    # Charger les chunks du client
    stmt = select(DocumentChunk).where(DocumentChunk.profile_id == profile_id)
    result = await db.execute(stmt)
    chunks = result.scalars().all()

    if not chunks:
        return []

    # Calculer la similarite cosinus
    scored = []
    for chunk in chunks:
        try:
            chunk_emb = json.loads(chunk.embedding)
            score = _cosine_similarity(query_embedding, chunk_emb)
            scored.append((chunk, score))
        except (json.JSONDecodeError, TypeError):
            continue

    # Trier par score decroissant
    scored.sort(key=lambda x: x[1], reverse=True)

    # Charger les documents parents pour le contexte
    results = []
    for chunk, score in scored[:top_k]:
        doc = await db.get(ClientDocument, chunk.document_id)
        if not doc:
            continue

        # Filtre theme optionnel
        if theme_filter and doc.themes:
            doc_themes = json.loads(doc.themes)
            if theme_filter.lower() not in [t.lower() for t in doc_themes]:
                continue

        results.append({
            "document_id": doc.id,
            "document_title": doc.title,
            "doc_type": doc.doc_type,
            "chunk_text": chunk.content,
            "chunk_idx": chunk.chunk_idx,
            "score": round(score, 4),
            "themes": json.loads(doc.themes) if doc.themes else [],
        })

    return results


def _cosine_similarity(a: list[float], b: list[float]) -> float:
    """Similarite cosinus entre deux vecteurs."""
    if len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(x * x for x in b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)
